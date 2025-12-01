"""
File service implementation.
Handles file operations using plug-and-play storage adapters.
Supports local filesystem, S3, and Supabase Storage.
"""
from pathlib import Path
from fastapi import UploadFile, HTTPException, status
from pypdf import PdfReader
import zipfile
import io
from typing import List, Tuple, Optional

# Import DOCX support (optional - will fail gracefully if not installed)
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. DOCX files will be stored but text extraction will fail.")

# Import DOC support (optional - requires textract library)
try:
    import textract
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False
    print("⚠️  textract not installed. DOC files will be stored but text extraction will fail. "
          "Install with: pip install textract (requires antiword or LibreOffice)")
from .interfaces import IFileService
from .storage import FileStorageFactory, FileStorageInterface
from ..core.config import STORAGE_TYPE, LOCAL_STORAGE_DIR, S3_BUCKET_NAME, AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_REGION, S3_ENDPOINT_URL, SUPABASE_URL, SUPABASE_KEY, SUPABASE_STORAGE_BUCKET

class FileService(IFileService):
    """
    File service implementation.
    Handles file operations - saving, reading, deleting files.
    Uses plug-and-play storage adapters (local, S3, Supabase).
    Follows Single Responsibility Principle.
    """
    
    def __init__(self, storage: FileStorageInterface = None):
        """
        Initialize file service with storage adapter.
        
        Args:
            storage: FileStorageInterface instance (if None, will be created from config)
        """
        self._storage: FileStorageInterface = storage
    
    async def get_storage(self) -> FileStorageInterface:
        """Get storage adapter (lazy initialization)."""
        if self._storage is None:
            self._storage = await self._create_storage()
        return self._storage
    
    async def _create_storage(self) -> FileStorageInterface:
        """Create storage adapter based on configuration."""
        if STORAGE_TYPE.lower() == "local":
            base_dir = Path(LOCAL_STORAGE_DIR) if LOCAL_STORAGE_DIR else None
            return await FileStorageFactory.create_and_initialize("local", base_dir=base_dir)
        elif STORAGE_TYPE.lower() == "s3":
            return await FileStorageFactory.create_and_initialize(
                "s3",
                bucket_name=S3_BUCKET_NAME,
                aws_access_key_id=AWS_ACCESS_KEY_ID,
                aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
                region_name=AWS_REGION,
                endpoint_url=S3_ENDPOINT_URL
            )
        elif STORAGE_TYPE.lower() == "supabase":
            return await FileStorageFactory.create_and_initialize(
                "supabase",
                supabase_url=SUPABASE_URL,
                supabase_key=SUPABASE_KEY,
                bucket_name=SUPABASE_STORAGE_BUCKET
            )
        else:
            raise ValueError(f"Unsupported STORAGE_TYPE: {STORAGE_TYPE}")
    
    async def save_upload(self, file: UploadFile, destination: Path) -> None:
        """
        Save uploaded file using storage adapter.
        
        Args:
            file: FastAPI UploadFile object
            destination: Path where file should be stored (relative to storage root)
        """
        try:
            storage_adapter = await self.get_storage()
            # Convert Path to string for storage adapter
            # If destination is absolute, extract relative path
            if destination.is_absolute():
                from ..core.config import UPLOAD_DIR
                try:
                    file_path = str(destination.relative_to(UPLOAD_DIR))
                except ValueError:
                    # If not under UPLOAD_DIR, use just the filename
                    file_path = destination.name
            else:
                file_path = str(destination)
            await storage_adapter.save_file(file, file_path)
        except Exception as e:
            print(f"Error saving upload: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Could not save file: {str(e)}"
            )
    
    async def extract_text(self, file_path: Path) -> str:
        """
        Extract text from file using storage adapter.
        
        Args:
            file_path: Storage path of the file (can be absolute or relative)
            
        Returns:
            Extracted text content
        """
        try:
            storage_adapter = await self.get_storage()
            
            # Handle both absolute and relative paths
            # If absolute path, convert to relative for storage adapter
            file_path_str = str(file_path)
            if file_path.is_absolute():
                # Try to extract relative path from uploads directory
                from ..core.config import UPLOAD_DIR
                try:
                    relative_path = file_path.relative_to(UPLOAD_DIR)
                    file_path_str = str(relative_path)
                except ValueError:
                    # If not under UPLOAD_DIR, use filename only
                    file_path_str = file_path.name
            
            # Check if file exists
            if not await storage_adapter.file_exists(file_path_str):
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"File not found for extraction: {file_path_str}"
                )
            
            # Get file content
            file_bytes = await storage_adapter.get_file(file_path_str)
            
            # Extract text based on file type
            file_ext = file_path.suffix.lower()
            
            if file_ext == ".pdf":
                # PDF text extraction
                try:
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                        tmp_file.write(file_bytes)
                        tmp_path = Path(tmp_file.name)
                    
                    try:
                        reader = PdfReader(tmp_path)
                        text_content = ""
                        for page in reader.pages:
                            text_content += page.extract_text() + "\n"
                        return text_content
                    finally:
                        # Clean up temp file
                        tmp_path.unlink()
                except Exception as e:
                    print(f"Error reading PDF: {e}")
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Error extracting text from PDF: {str(e)}"
                    )
            
            elif file_ext == ".docx":
                # DOCX text extraction using python-docx
                if not DOCX_AVAILABLE:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="DOCX support not available. Please install python-docx: pip install python-docx"
                    )
                
                try:
                    doc = DocxDocument(io.BytesIO(file_bytes))
                    text_content = ""
                    
                    # Extract text from paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content += paragraph.text + "\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_content += " | ".join(row_text) + "\n"
                    
                    if not text_content.strip():
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail="DOCX file appears to be empty or contains no extractable text"
                        )
                    
                    return text_content
                except Exception as e:
                    print(f"Error reading DOCX: {e}")
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Error extracting text from DOCX: {str(e)}"
                    )
            
            elif file_ext == ".doc":
                # DOC (old Word format) - requires textract library
                if not DOC_AVAILABLE:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="DOC (old Word format) extraction requires 'textract' library. "
                               "Install with: pip install textract (requires antiword or LibreOffice). "
                               "Alternatively, convert DOC to DOCX before uploading."
                    )
                
                try:
                    # Use textract to extract text from DOC files
                    # textract requires system dependencies (antiword or LibreOffice)
                    text_content = textract.process(io.BytesIO(file_bytes), extension='doc').decode('utf-8')
                    
                    if not text_content.strip():
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail="DOC file appears to be empty or contains no extractable text"
                        )
                    
                    return text_content
                except HTTPException:
                    raise
                except Exception as e:
                    print(f"Error reading DOC: {e}")
                    error_msg = str(e)
                    if "antiword" in error_msg.lower() or "libreoffice" in error_msg.lower():
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail=f"DOC extraction failed: {error_msg}. "
                                   "Please ensure antiword or LibreOffice is installed on the system. "
                                   "Alternatively, convert DOC to DOCX before uploading."
                        )
                    else:
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail=f"Error extracting text from DOC: {error_msg}"
                        )
            
            elif file_ext in [".txt", ".md", ".markdown"]:
                # Plain text files (TXT, Markdown)
                try:
                    return file_bytes.decode('utf-8', errors='ignore')
                except Exception as e:
                    print(f"Error reading text file: {e}")
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Error reading text file: {str(e)}"
                    )
            
            else:
                # Unknown file type - try UTF-8 decoding as fallback
                try:
                    decoded = file_bytes.decode('utf-8', errors='ignore')
                    if decoded.strip():
                        return decoded
                    else:
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail=f"File format '{file_ext}' is not supported for text extraction. "
                                   "Supported formats: PDF, DOCX, DOC, TXT, MD"
                        )
                except Exception as e:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"File format '{file_ext}' is not supported for text extraction. "
                               "Supported formats: PDF, DOCX, DOC, TXT, MD. Error: {str(e)}"
                    )
        except HTTPException:
            raise
        except Exception as e:
            print(f"General file error: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error processing file: {str(e)}"
            )
    
    async def delete_file(self, file_path: Path) -> None:
        """
        Delete a file using storage adapter.
        
        Args:
            file_path: Storage path of the file to delete
        """
        try:
            storage_adapter = await self.get_storage()
            file_path_str = str(file_path)
            await storage_adapter.delete_file(file_path_str)
        except Exception as e:
            print(f"Error deleting file {file_path}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Could not delete file: {str(e)}"
            )
    
    async def save_markdown(self, content: str, destination: Path) -> None:
        """
        Save markdown content using storage adapter.
        
        Args:
            content: Markdown content to save
            destination: Storage path where content should be saved (can be absolute or relative)
        """
        try:
            storage_adapter = await self.get_storage()
            # Convert absolute path to relative path for storage adapter
            from ..core.config import UPLOAD_DIR
            if destination.is_absolute():
                try:
                    # Try to get relative path from UPLOAD_DIR
                    relative_path = destination.relative_to(UPLOAD_DIR)
                    file_path = str(relative_path)
                except ValueError:
                    # If not under UPLOAD_DIR, use filename only
                    file_path = destination.name
            else:
                file_path = str(destination)
            await storage_adapter.save_text(content, file_path)
        except Exception as e:
            print(f"Error saving markdown: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Could not save markdown: {str(e)}"
            )
    
    async def get_file_url(self, file_path: Path, expires_in: int = 3600) -> str:
        """
        Get a URL to access the file (for direct download/viewing).
        
        Args:
            file_path: Storage path of the file
            expires_in: URL expiration time in seconds (for signed URLs)
        
        Returns:
            URL string for accessing the file
        """
        storage_adapter = await self.get_storage()
        file_path_str = str(file_path)
        return await storage_adapter.get_file_url(file_path_str, expires_in=expires_in)
    
    async def extract_zip_files(self, zip_file: UploadFile, base_folder: Optional[str] = None) -> Tuple[List[Tuple[bytes, str, Optional[str]]], List[str]]:
        """
        Extract files from a ZIP archive.
        
        Args:
            zip_file: UploadFile containing ZIP archive
            base_folder: Optional base folder path to prepend to extracted file paths
        
        Returns:
            Tuple of:
            - List of tuples: (file_bytes, filename, folder_path)
              - file_bytes: Content of the extracted file
              - filename: Name of the file
              - folder_path: Folder path within ZIP (None for root files)
            - List of folder paths found in ZIP (including empty folders)
        
        Raises:
            HTTPException: If ZIP file is invalid or corrupted
        """
        try:
            # Read ZIP file content
            zip_content = await zip_file.read()
            
            # Validate ZIP file
            if not zipfile.is_zipfile(io.BytesIO(zip_content)):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid ZIP file format"
                )
            
            extracted_files = []
            folder_paths = set()  # Track all folder paths (including empty ones)
            
            with zipfile.ZipFile(io.BytesIO(zip_content), 'r') as zip_ref:
                # Get list of files in ZIP
                file_list = zip_ref.namelist()
                
                for file_path in file_list:
                    # Skip hidden files and system files
                    if file_path.startswith('__MACOSX/') or file_path.startswith('.DS_Store'):
                        continue
                    
                    # Track folder paths (including empty directories)
                    if file_path.endswith('/'):
                        # This is a directory entry
                        folder_path_clean = file_path.rstrip('/')
                        if folder_path_clean:
                            # Build full folder path
                            if base_folder:
                                full_folder_path = f"{base_folder}/{folder_path_clean}" if folder_path_clean else base_folder
                            else:
                                full_folder_path = folder_path_clean if folder_path_clean else None
                            
                            if full_folder_path:
                                # Add all parent folders too
                                parts = full_folder_path.split('/')
                                for i in range(1, len(parts) + 1):
                                    parent_folder = '/'.join(parts[:i])
                                    folder_paths.add(parent_folder)
                        continue
                    
                    try:
                        # Extract file content
                        file_bytes = zip_ref.read(file_path)
                        
                        # Extract filename and folder path
                        path_parts = file_path.split('/')
                        filename = path_parts[-1]
                        
                        # Build folder path (everything except filename)
                        if len(path_parts) > 1:
                            zip_folder = '/'.join(path_parts[:-1])
                            # Combine with base folder if provided
                            if base_folder:
                                folder_path = f"{base_folder}/{zip_folder}" if zip_folder else base_folder
                            else:
                                folder_path = zip_folder if zip_folder else None
                            
                            # Add all parent folders
                            if folder_path:
                                parts = folder_path.split('/')
                                for i in range(1, len(parts) + 1):
                                    parent_folder = '/'.join(parts[:i])
                                    folder_paths.add(parent_folder)
                        else:
                            folder_path = base_folder
                        
                        # Only process supported file types
                        # ✅ Full support (text extraction): PDF, TXT, MD, DOCX, DOC
                        # ⚠️ Storage only (no extraction): RTF, ODT
                        file_ext = Path(filename).suffix.lower()
                        supported_extensions = {'.pdf', '.txt', '.md', '.docx', '.doc', '.rtf', '.odt'}
                        
                        if file_ext in supported_extensions or not file_ext:
                            extracted_files.append((file_bytes, filename, folder_path))
                        else:
                            print(f"Skipping unsupported file type: {filename} ({file_ext})")
                    
                    except Exception as e:
                        print(f"Error extracting file {file_path} from ZIP: {e}")
                        continue
            
            if not extracted_files and not folder_paths:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="ZIP file contains no supported files or folders"
                )
            
            return extracted_files, sorted(list(folder_paths))
        
        except HTTPException:
            raise
        except zipfile.BadZipFile:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid or corrupted ZIP file"
            )
        except Exception as e:
            print(f"Error extracting ZIP file: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error extracting ZIP file: {str(e)}"
            )
