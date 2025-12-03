"""
DOCX Text Extractor.

Extracts text from DOCX files using python-docx library.
"""
import io
from .base import BaseTextExtractor
from fastapi import HTTPException, status
from ...core.logging_config import get_logger

logger = get_logger(__name__)

# Check if python-docx is available
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX files will be stored but text extraction will fail.")


class DOCXExtractor(BaseTextExtractor):
    """Extractor for DOCX files."""
    
    def __init__(self):
        super().__init__(".docx", "DOCX")
    
    def _check_availability(self) -> bool:
        """Check if python-docx is installed."""
        return DOCX_AVAILABLE
    
    def get_error_message(self) -> str:
        """Get error message when DOCX support is not available."""
        return "DOCX support not available. Please install python-docx: pip install python-docx"
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If DOCX extraction fails or library not available
        """
        if not self.is_available():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=self.get_error_message()
            )
        
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            # Validate content
            self.validate_content(text_content)
            
            return text_content
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Error extracting text from DOCX: {str(e)}"
            )

